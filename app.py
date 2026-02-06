import streamlit as st
import os
from dotenv import load_dotenv
import google.generativeai as genai
from pypdf import PdfReader
import json
from datetime import datetime
import io
import time
import random
import streamlit.components.v1 as components

# --- 1. SEGURIDAD DE LIBRERÍAS (INTACTO) ---
try:
    from docx import Document
    from docx.shared import Pt
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# --- 2. CONFIGURACIÓN (INTACTO) ---
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

if not api_key:
    st.error("⚠️ Error: No se encuentra la API Key en el archivo .env")
    st.stop()

genai.configure(api_key=api_key)

st.set_page_config(
    page_title="GodzillaBot Oposiciones", 
    page_icon="🦖", 
    layout="wide",
    initial_sidebar_state="collapsed" 
)

DOCS_DIR = "documentos"
HISTORY_DIR = "historial_sesiones"
if not os.path.exists(DOCS_DIR): os.makedirs(DOCS_DIR)
if not os.path.exists(HISTORY_DIR): os.makedirs(HISTORY_DIR)

# --- 3. GESTIÓN DE ESTADO (INTACTO) ---
if "pdf_text" not in st.session_state:
    st.session_state.pdf_text = ""
if "last_files" not in st.session_state:
    st.session_state.last_files = []

# --- 4. LÓGICA (INTACTO) ---
@st.cache_data(show_spinner=False)
def get_pdf_text_fast(file_names):
    text = ""
    for name in file_names:
        try:
            path = os.path.join(DOCS_DIR, name)
            reader = PdfReader(path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except: continue
    return text

@st.cache_resource
def get_model_list():
    try:
        all_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        flash_models = [m for m in all_models if 'flash' in m.lower()]
        pro_models = [m for m in all_models if 'pro' in m.lower() and 'vision' not in m.lower()]
        priorities = flash_models + pro_models
        if not priorities: return ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
        return priorities
    except:
        return ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']

MODELS_AVAILABLE = get_model_list()

def generate_response_with_patience(prompt_text):
    max_retries = 3
    frases_espera = [
        "🦖 Godzilla está recargando su aliento atómico...",
        "⏳ Negociando prioridad con Google...",
        "🦕 Masticando gigas de normativa...",
        "🔥 Buscando el artículo exacto en la red neuronal..."
    ]
    
    for attempt in range(max_retries):
        for model_name in MODELS_AVAILABLE:
            try:
                model = genai.GenerativeModel(model_name)
                return model.generate_content(prompt_text, stream=True)
            except Exception as e:
                error_msg = str(e)
                if "429" in error_msg or "quota" in error_msg.lower():
                    wait_time = (attempt + 1) * 2 
                    if attempt >= 0: 
                         msg = random.choice(frases_espera)
                         st.toast(f"{msg} (Reintentando...)", icon="🦖")
                    time.sleep(wait_time)
                    continue
                if "404" in error_msg:
                    continue
                continue
    return "Error_Quota_Final"

def auto_scroll():
    js = """
    <script>
        setTimeout(function() {
            var body = window.parent.document.querySelector(".main");
            if (body) { body.scrollTop = body.scrollHeight; }
        }, 300);
    </script>
    """
    components.html(js, height=0, width=0)

# --- 5. ESTÉTICA OPTIMIZADA (CSS MODIFICADO PARA MENÚ) ---
st.markdown("""
<style>
    /* 1. CONFIGURACIÓN BASE */
    [data-testid="stHeader"] { 
        background-color: transparent !important; 
        z-index: 1 !important; 
        height: 60px !important; /* Asegurar altura para que no oculte cosas */
    }
    
    /* Ocultar elementos molestos, PERO NO EL CONTENEDOR DEL MENÚ */
    [data-testid="stToolbar"] { display: none !important; }
    [data-testid="stDecoration"] { display: none !important; }
    footer { visibility: hidden; }

    /* 2. FONDO Y TEXTO */
    .stApp, [data-testid="stAppViewContainer"], .main { 
        background-color: #ffffff !important; 
        font-family: 'Segoe UI', Helvetica, Arial, sans-serif;
    }

    /* 3. ¡¡¡EL BOTÓN DEL MENÚ (LA CLAVE)!!! */
    /* Lo forzamos a ser visible, fijo y por encima de todo */
    [data-testid="stSidebarCollapsedControl"] {
        display: flex !important;
        visibility: visible !important;
        background-color: white !important;
        color: #16a34a !important; /* Verde */
        border: 2px solid #16a34a !important;
        border-radius: 50% !important;
        width: 50px !important;
        height: 50px !important;
        align-items: center;
        justify-content: center;
        
        /* POSICIONAMIENTO FIJO AGRESIVO */
        position: fixed !important;
        top: 10px !important;
        left: 10px !important;
        z-index: 9999999 !important; /* Por encima del propio Dios */
        
        box-shadow: 0 4px 10px rgba(0,0,0,0.2) !important;
    }
    
    /* Asegurar que el icono dentro del botón sea verde */
    [data-testid="stSidebarCollapsedControl"] svg {
        fill: #16a34a !important;
        width: 30px !important;
        height: 30px !important;
    }

    /* 4. CHAT BUBBLES */
    .stChatMessage { background-color: transparent !important; }
    div[data-testid="stChatMessage"]:nth-child(odd) { 
        background-color: #f3f4f6 !important; 
        border: 1px solid #e5e7eb !important; border-radius: 12px !important;
        color: #111827 !important; padding: 12px !important;
    }
    div[data-testid="stChatMessage"]:nth-child(even) { 
        background-color: #ffffff !important; 
        border: 1px solid #d1d5db !important; border-radius: 12px !important;
        color: #000000 !important; padding: 12px !important;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05); 
    }
    div[data-testid="stChatMessage"] * { color: inherit !important; }

    /* 5. INPUT IPHONE FIX */
    .stChatInput textarea { font-size: 16px !important; }

    /* 6. CABECERA */
    .header-container {
        background: linear-gradient(90deg, #166534 0%, #15803d 100%);
        padding: 20px; border-radius: 10px; color: white; text-align: center;
        margin-bottom: 20px; 
        margin-top: 50px; /* Margen superior para que el botón no tape el título */
        box-shadow: 0 4px 6px rgba(0,0,0,0.1); border-bottom: 4px solid #4ade80;
    }
    .header-container h1 { font-size: 2.2rem; margin: 0; font-weight: 800;}
    .header-container p { font-size: 1rem; opacity: 0.9; margin-top: 5px; font-style: italic; }

    /* 7. MÓVIL RESPONSIVE */
    @media only screen and (max-width: 768px) {
        .block-container { padding-top: 4rem !important; padding-left: 0.5rem; padding-right: 0.5rem; }
        .header-container { padding: 15px; margin-top: 45px; }
        .header-container h1 { font-size: 1.6rem !important; }
    }

    /* 8. LANDSCAPE */
    @media only screen and (orientation: landscape) and (max-height: 600px) {
        .block-container { padding-top: 0.5rem !important; }
        .header-container {
            padding: 5px !important; margin-bottom: 10px !important; margin-top: 0px !important;
            display: flex; align-items: center; justify-content: center; min-height: 40px;
        }
        .header-container h1 { font-size: 1.1rem !important; margin: 0; }
        .header-container p { display: none !important; }
        /* Botón más pequeño en horizontal */
        [data-testid="stSidebarCollapsedControl"] {
            width: 35px !important; height: 35px !important; top: 5px !important; left: 5px !important;
        }
    }
    
    /* 9. TABLAS */
    div[data-testid="stMarkdownContainer"] table { width: 100%; border-collapse: collapse !important; border: 1px solid #374151 !important; }
    div[data-testid="stMarkdownContainer"] th { background-color: #e5e7eb !important; color: #000000 !important; border: 1px solid #9ca3af !important; padding: 6px; }
    div[data-testid="stMarkdownContainer"] td { border: 1px solid #d1d5db !important; padding: 6px; color: #000000; vertical-align: top; }
</style>
""", unsafe_allow_html=True)

# --- 6. FUNCIONES DOCS (INTACTO) ---
def create_word_docx(text_content):
    if not WORD_AVAILABLE: return None
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(11)
    doc.add_heading('Informe GodzillaBot', 0)
    for line in text_content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('**') and line.endswith('**'): 
            p = doc.add_paragraph(); p.add_run(line.replace('**', '')).bold = True
        else:
            doc.add_paragraph(line.replace('**', ''))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def save_uploaded_file(uploaded_file):
    try:
        with open(os.path.join(DOCS_DIR, uploaded_file.name), "wb") as f:
            f.write(uploaded_file.getbuffer())
        return True
    except: return False

def save_session_history():
    if not st.session_state.messages: return
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    path = os.path.join(HISTORY_DIR, f"Sesion_{timestamp}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(st.session_state.messages, f, ensure_ascii=False, indent=4)
    st.success(f"✅")

def load_session_history(filename):
    path = os.path.join(HISTORY_DIR, filename)
    try:
        with open(path, "r", encoding="utf-8") as f:
            st.session_state.messages = json.load(f)
        st.rerun()
    except: st.error("Error")

def get_system_prompt(mode):
    base = "Eres GodzillaBot, experto en legislación. Fuente: PDFs adjuntos. "
    if "Simulacro" in mode:
        pers = ["EL MINUCIOSO", "EL CONCEPTUAL", "EL TRAMPOSO", "EL PRÁCTICO"]
        examinador = random.choice(pers)
        return base + f"""
        MODO: SIMULACRO TEST. Personalidad: {examinador}.
        INSTRUCCIONES DE FORMATO:
        1. Opciones (a, b, c, d) en LÍNEAS SEPARADAS (Lista Markdown).
        2. Hoja de Respuestas al final.
        """
    elif "Excel" in mode:
        return base + "Salida: Tabla Markdown cerrada. Concepto | Dato | Art | Nota."
    else:
        return base + "Responde de forma técnica y estructurada."

# --- 7. MENÚ LATERAL (INTACTO) ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/1624/1624022.png", width=70) 
    st.markdown("## 🦖 Guarida")
    
    with st.expander("📤 Cargar Temario (PDFs)", expanded=False):
        up = st.file_uploader("Arrastra archivos aquí", type="pdf")
        if up and save_uploaded_file(up): st.rerun()
    
    files_available = [f for f in os.listdir(DOCS_DIR) if f.endswith('.pdf')]
    if files_available:
        files = st.multiselect("📚 Documentos Activos:", files_available, default=[]) 
    else:
        files = []
        st.info("ℹ️ Sube PDFs para empezar.")
    
    st.markdown("---")
    st.markdown("### 🎯 Objetivo de Hoy")
    mode = st.radio("Estrategia:", 
        ["💀 Simulacro de Examen (Test)", "💬 Chat Interactivo con Temario", 
         "📝 Resumen de Alto Rendimiento", "📊 Extracción de Datos a Excel"])
    
    st.markdown("---")
    c1, c2 = st.columns(2)
    if c1.button("💾 Guardar"): save_session_history()
    if c2.button("🗑️ Reiniciar"): st.session_state.messages = []; st.rerun()
    
    sessions = [f for f in os.listdir(HISTORY_DIR) if f.endswith('.json')]
    if sessions:
        load = st.selectbox("Recuperar:", ["..."] + sorted(sessions, reverse=True))
        if load != "..." and st.button("Abrir"): load_session_history(load)

# --- 8. ZONA PRINCIPAL (INTACTO) ---
st.markdown("""
<div class="header-container">
    <h1>🦖 GodzillaBot Oposiciones</h1>
    <p>Destruyendo tus dudas, dominando el temario.</p>
</div>
""", unsafe_allow_html=True)

if files != st.session_state.last_files:
    if files:
        with st.spinner("🦖 Digiriendo documentos nuevos..."):
            st.session_state.pdf_text = get_pdf_text_fast(files)
            st.session_state.last_files = files
    else:
        st.session_state.pdf_text = ""
        st.session_state.last_files = []

if "messages" not in st.session_state: st.session_state.messages = []

for i, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]): 
        st.markdown(msg["content"])
        if msg["role"] == "assistant":
            key_base = f"btn_{i}"
            col1, col2 = st.columns([1, 1])
            with col1:
                if WORD_AVAILABLE:
                    docx = create_word_docx(msg["content"])
                    st.download_button("📄 Word", docx, f"Godzilla_{i}.docx", 
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"{key_base}_w")
            with col2:
                if "|" in msg["content"]:
                    st.download_button("📊 Excel", msg["content"], f"datos_{i}.csv", "text/csv", key=f"{key_base}_x")

if prompt := st.chat_input("Escribe tu pregunta..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"): st.markdown(prompt)
    auto_scroll()

    if not files:
        st.warning("⚠️ Selecciona documentos en el menú lateral.")
    else:
        with st.chat_message("assistant"):
            placeholder = st.empty()
            full_resp = ""
            try:
                text_context = st.session_state.pdf_text 
                
                with st.spinner("🦖 Godzilla responde..."): 
                    prompt_final = f"{get_system_prompt(mode)}\nDOCS: {text_context[:800000]}\nUSER: {prompt}"
                    response_obj = generate_response_with_patience(prompt_final)

                    if isinstance(response_obj, str) and response_obj.startswith("Error_Quota"):
                        st.error("🛑 Agotado total. Necesito 2 min.")
                    else:
                        for chunk in response_obj:
                            if chunk.text:
                                full_resp += chunk.text
                                placeholder.markdown(full_resp + "▌")
                        placeholder.markdown(full_resp)
                        st.session_state.messages.append({"role": "assistant", "content": full_resp})
                        st.rerun() 
            except Exception as e: st.error(f"Error: {e}")